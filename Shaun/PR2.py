from PIL import Image, ImageEnhance, ImageOps
from docx import Document
from difflib import SequenceMatcher
from transformers import pipeline
import os
import cv2
import numpy as np
import easyocr

# Initialize AI Model for Semantic Comparison
similarity_model = pipeline("text-classification", model="cross-encoder/nli-distilroberta-base")
reader = easyocr.Reader(['en'])  # Initialize EasyOCR for English text recognition

def preprocess_image(image_path):
    """Preprocess the image for better OCR results."""
    try:
        image = cv2.imread(image_path, cv2.IMREAD_GRAYSCALE)
        # Increase contrast and remove noise
        image = cv2.GaussianBlur(image, (5, 5), 0)
        _, image = cv2.threshold(image, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        preprocessed_path = f"{image_path}_preprocessed.png"
        cv2.imwrite(preprocessed_path, image)
        return preprocessed_path
    except Exception as e:
        print(f"Error preprocessing image: {e}")
        return image_path

def extract_text_from_image(image_path):
    """Extract text from an image using EasyOCR."""
    try:
        preprocessed_path = preprocess_image(image_path)
        result = reader.readtext(preprocessed_path)
        os.remove(preprocessed_path)  # Clean up preprocessed image
        text = " ".join([item[1] for item in result]).strip()
        return text
    except Exception as e:
        return f"Error: Could not process image - {e}"

def calculate_similarity(text1, text2):
    """Calculate similarity using Levenshtein distance and semantic analysis."""
    # Levenshtein-based similarity
    ratio = SequenceMatcher(None, text1, text2).ratio()
    # Semantic similarity using AI model
    semantic_result = similarity_model(f"{text1} [SEP] {text2}")
    semantic_score = semantic_result[0]['score']
    return max(ratio, semantic_score)  # Return the higher of the two scores

def extract_images_and_text(doc_path, temp_folder):
    """Extract images and their corresponding scanned text from the Word file."""
    document = Document(doc_path)
    images_and_text = []
    counter = 0

    for rel in document.part.rels.values():
        if "image" in rel.target_ref:
            # Save image to temporary folder
            image_path = os.path.join(temp_folder, f"image_{counter}.png")
            with open(image_path, "wb") as img_file:
                img_file.write(rel.target_part.blob)
            # Get the scanned text (next paragraph)
            if counter < len(document.paragraphs):
                scanned_text = document.paragraphs[counter].text.strip()
                images_and_text.append((image_path, scanned_text))
            counter += 1
    return images_and_text

def create_error_report(errors, output_path):
    """Create a Word document with the errors."""
    output_doc = Document()
    output_doc.add_heading("Proofreading Error Report", level=1)

    for error in errors:
        image_name, extracted_text, scanned_text, similarity = error
        output_doc.add_heading(f"Image: {image_name}", level=2)
        output_doc.add_paragraph(f"Extracted Text:\n{extracted_text}")
        output_doc.add_paragraph(f"Scanned Text:\n{scanned_text}")
        output_doc.add_paragraph(f"Similarity Score: {similarity:.2f}")
        output_doc.add_paragraph("-" * 40)

    if not errors:
        output_doc.add_paragraph("No mismatches found. All text matches perfectly!")
    
    output_doc.save(output_path)

def proofread_document_with_ai(doc_path, output_path, temp_folder="temp_images", similarity_threshold=0.8):
    """Proofread the Word document and create an error report."""
    if not os.path.exists(temp_folder):
        os.makedirs(temp_folder)
    
    images_and_text = extract_images_and_text(doc_path, temp_folder)
    errors = []

    for image_path, scanned_text in images_and_text:
        # Extract text from the image
        extracted_text = extract_text_from_image(image_path)
        # Calculate similarity
        similarity = calculate_similarity(extracted_text, scanned_text)
        # Log errors if similarity is below the threshold
        if similarity < similarity_threshold:
            errors.append((os.path.basename(image_path), extracted_text, scanned_text, similarity))
        os.remove(image_path)  # Clean up temporary image

    # Create an error report
    create_error_report(errors, output_path)
    print(f"Proofreading complete. Errors saved to: {output_path}")

# Paths
input_doc_path = "/workspaces/first/Shaun/remem3.docx"  # Your input Word file
output_doc_path = "/workspaces/first/Shaun/ai_errors_report.docx"  # Output file for the error report

# Run the proofreading function
proofread_document_with_ai(input_doc_path, output_doc_path)
