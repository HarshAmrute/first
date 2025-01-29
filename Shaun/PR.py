import pytesseract
from PIL import Image
from docx import Document
import os

# Set up Tesseract path if necessary (e.g., for Windows users)
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

def extract_text_from_image(image_path):
    """Extract text from an image using Tesseract OCR."""
    try:
        image = Image.open(image_path)
        return pytesseract.image_to_string(image).strip()
    except Exception as e:
        print(f"Error processing image {image_path}: {e}")
        return ""

def proofread_document(doc_path, output_path):
    """Proofread and correct text in a Word document."""
    # Open the Word document
    document = Document(doc_path)
    new_document = Document()
    
    paragraphs = document.paragraphs
    for i, paragraph in enumerate(paragraphs):
        if paragraph.runs and paragraph.runs[0].element.tag.endswith('drawing'):  # Detect if it's an image
            # Save the image temporarily for OCR processing
            image_name = f"temp_image_{i}.png"
            with open(image_name, "wb") as f:
                f.write(paragraph.runs[0].element.blob)
            
            # Extract text from the image
            extracted_text = extract_text_from_image(image_name)
            os.remove(image_name)  # Remove temporary image
            
            # Add the extracted text to the new document
            new_document.add_paragraph(f"<Image OCR Text>: {extracted_text}")
        
        else:
            # Copy text as is to the new document
            new_document.add_paragraph(paragraph.text)
    
    # Save the corrected document
    new_document.save(output_path)
    print(f"Proofread document saved at: {output_path}")

# Paths to your Word file and output file
input_doc_path = "/workspaces/first/Shaun/remem3.docx"  # Replace with your Word file path
output_doc_path = "output_corrected.docx"

proofread_document(input_doc_path, output_doc_path)
