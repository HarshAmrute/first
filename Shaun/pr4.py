import cv2
import numpy as np
from docx import Document
import os

# Function to preprocess image
def preprocess_image(image_path):
    """Preprocess the image to enhance text for OCR extraction."""
    # Load image
    img = cv2.imread(image_path)
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)  # Convert to grayscale
    
    # Enhance image with thresholding and denoise
    _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY_INV)  # Inverted binary threshold
    denoised = cv2.fastNlMeansDenoising(thresh, None, 30, 7, 21)  # Denoising
    
    # Save the processed image for further steps
    processed_path = image_path.replace(".png", "_processed.png")
    cv2.imwrite(processed_path, denoised)
    return processed_path

def extract_text_with_custom_model(image_path):
    """Use OpenCV's preprocessing with a custom model to extract text (for the sake of example)."""
    preprocessed_image = preprocess_image(image_path)
    
    # Here you can integrate a custom trained OCR model or use any open model you may have.
    # This part is where you'd integrate an OCR model.
    # For now, we're assuming this function returns dummy text after OCR preprocessing.
    return "Extracted Text Placeholder"

def extract_images_from_docx(doc_path, temp_folder="temp_images"):
    """Extract images from Word document."""
    document = Document(doc_path)
    if not os.path.exists(temp_folder):
        os.makedirs(temp_folder)

    image_paths = []
    for rel in document.part.rels.values():
        if "image" in rel.target_ref:
            image_path = os.path.join(temp_folder, f"{os.path.basename(rel.target_ref)}.png")
            with open(image_path, "wb") as img_file:
                img_file.write(rel.target_part.blob)
            image_paths.append(image_path)
    
    return image_paths

def create_word_document_with_extracted_text(doc_path, output_doc_path, temp_folder="temp_images"):
    """Process the Word document and add OCR text back to a new document."""
    image_paths = extract_images_from_docx(doc_path, temp_folder)
    
    # Create a new Word document
    new_doc = Document()
    new_doc.add_heading('Extracted Text from Images', 0)

    for image_path in image_paths:
        extracted_text = extract_text_with_custom_model(image_path)
        new_doc.add_heading(f"Text from {os.path.basename(image_path)}", level=2)
        new_doc.add_paragraph(extracted_text)
        
    # Save the new document with extracted text
    new_doc.save(output_doc_path)
    print(f"Text extraction complete. The new document is saved at: {output_doc_path}")

# Your specific file paths
input_doc_path = "/workspaces/first/Shaun/remem3.docx"  # Update this to your actual input doc path
output_doc_path = "/workspaces/first/Shaun/output_with_extracted_text.docx"  # Path where output doc should be saved

# Run the function to extract text from images and create a new Word document
create_word_document_with_extracted_text(input_doc_path, output_doc_path)
