import re
from docx import Document

def extract_content_from_docx(doc_path):
    """Extract content enclosed in < > and " " (including smart quotes) from a Word document."""
    # Load the document
    doc = Document(doc_path)
    
    extracted_content = []

    # Define regex patterns to match content inside <> and "" (both regular and smart quotes)
    pattern_angle_brackets = r'<[^>]+>'  # Matches content inside < >
    pattern_quotes = r'[“”""]([^“”""]+)[“”""]'  # Matches content inside “ ” or " "

    # Iterate over each paragraph in the document
    for para in doc.paragraphs:
        # Find all occurrences of content inside angle brackets
        angle_bracket_matches = re.findall(pattern_angle_brackets, para.text)
        extracted_content.extend(angle_bracket_matches)

        # Find all occurrences of content inside double quotes (including smart quotes)
        quote_matches = re.findall(pattern_quotes, para.text)
        # Retain the quotes with the extracted content
        for match in quote_matches:
            # Ensure the quotes are kept in the extracted content
            extracted_content.append(f'“{match}”')  # Adding smart quotes back

    return extracted_content

def create_reduced_docx(extracted_content, output_path):
    """Create a new Word document with only the extracted content."""
    # Create a new Document
    new_doc = Document()
    
    # Add extracted content to the new document
    for content in extracted_content:
        new_doc.add_paragraph(content)
    
    # Save the new document
    new_doc.save(output_path)

# Path to the input Word document
input_doc_path = "/workspaces/first/Shaun/remem3.docx"  # Your input Word file path
output_doc_path = "/workspaces/first/Shaun/reduced_content.docx"  # The output reduced docx file path

# Extract content from the original document
extracted_content = extract_content_from_docx(input_doc_path)

# Create a new Word document with the extracted content
create_reduced_docx(extracted_content, output_doc_path)

print(f"Reduced document saved as {output_doc_path}")
